import re
import itertools
import threading
import time
import sys

##extra libraries to install
import xlrd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
import easygui



# stname1 = input('enter placement status file name');
# stname2 = input('enter cse student status file name');
# stname1 = '1.xlsx'
# stname2 = '2.xlsx'
stname1 = easygui.fileopenbox("select placement status file name")
stname2 = easygui.fileopenbox("select cse student status file name")


done = False
global stt
# stt = '\rloading '
#here is the animation
def animate():
    count=1

    stt = '\rProcessing'
    for c in itertools.cycle(['|', '/', '-', '\\']):
        if done:
            break
        stt += '*'
        sys.stdout.write(stt + c)
        sys.stdout.flush()
        time.sleep(0.4)
    sys.stdout.write('\rDone!     \n\n')

t = threading.Thread(target=animate)



# shutil.copy('empty.docx','companywise_ug.docx')
# shutil.copy('empty.docx','companywise_pg.docx')
# shutil.copy('empty.docx','no_of_offers_ug.docx')
# shutil.copy('empty.docx','no_of_offers_pg.docx')
# shutil.copy('empty.docx','student_ug.docx')
# shutil.copy('empty.docx','student_pg.docx')


def closest(commands, u):
    return min(commands, key=lambda v: len(set(u) ^ set(v)))


salary_temp_copy = []


def salary_peocessing(salary_temp):
    count = 0
    for i in range(salary_temp.__len__()):
        if (salary_temp[i] == ''):
            salary_temp[i] = 'NULL'



    for i in range(salary_temp.__len__()):
        st = str(salary_temp[i])

        if (st != 'NULL'):
            st = st.upper()
            st = st.replace('RS', '')
            st = st.replace(',', '')
            st = st.replace('STIPEND', '')
            st = st.replace('(', '')
            st = st.replace(')', '')
            st = st.replace('UG', 'UG ')
            st = st.replace('PG', 'PG ')
            st = st.replace('MCA', '')
            st = st.replace('/', '')
            st = st.replace('\\', '')
            st = st.replace('K', '000')
            st = st.replace(' K', '000')
            x = st.split(" ")

            st2 = st.replace(' ','')
            st2 = st2.replace('.', '')

            while '' in x:
                x.remove('')




            if 'UG' in x or 'PG' in x:
                if 'LPA' in x:
                    for k3 in range(x.__len__()):
                        if (x[k3] == 'UG'):
                            while (1):
                                if(k3 < x.__len__()-1):
                                    k3 += 1
                                    if (x[k3] == 'LPA'):
                                        st = x[k3 - 1]
                                        break
                                else:
                                    break


                    for k3 in range(x.__len__()):
                        if (x[k3] == 'PG'):
                            while (1):
                                if (k3 < x.__len__() -1):
                                    k3 += 1
                                    if (x[k3] == 'LPA'):
                                        st = st + '/' + x[k3 - 1]
                                        break
                                else:
                                    break


                else:
                    if 'PM' in x:
                        for k3 in range(x.__len__()):
                            if (x[k3] == 'UG'):
                                while (1):
                                    if (k3 < x.__len__()-1):
                                        k3 += 1
                                        if (x[k3] == 'PM'):
                                            st = x[k3 - 1]
                                            num = int(st)
                                            num = num * 12
                                            num = num / 100000
                                            st = str(round(num, 1))
                                            break
                                    else:
                                        break

                        for k3 in range(x.__len__()):
                            if (x[k3] == 'PG'):
                                while (1):
                                    if (k3 < x.__len__()-1):
                                        k3 += 1
                                        if (x[k3] == 'PM'):
                                            st1 = x[k3 - 1]
                                            num = int(st1)
                                            num = num * 12
                                            num = num / 100000
                                            st1 = str(round(num, 1))
                                            st = st + '/' + st1

                                            break

                                    else:
                                        break

                pass
            else:
                if 'LPA' in x:
                    for k1 in range(x.__len__()):
                        if (x[k1] == 'LPA'):
                            st = x[k1 - 1]
                else:
                    if 'PM' in x:
                        for k2 in range(x.__len__()):
                            if (x[k2] == 'PM'):
                                st = x[k2 - 1]
                                num = int(st)
                                num = num * 12
                                num = num / 100000
                                st = str(round(num, 1))

            salary_temp[i] = st
        else:
            st3 = str(comp_info[i]).upper()
            if not st3.__contains__('TOTAL'):
                sl = 'enter salary for comapny '+ str(comp_info[i])
                sal2 = float(input(sl))
                salary_temp[i] = sal2


    return salary_temp


def salary_check(salary_temp):
    for i in range(salary_temp.__len__()):
        s = str(salary_temp[i])
        s = s.replace('.', '')
        s = s.replace('/', '')
        if (((re.search('\D', s)) and s != '0'and s!=0 and s != 'NULL') or s.__len__() > 8):
            print(salary_temp[i])
            # print('It conatins non-readable string please convert it to standard form')
            # print('ex: if string has \'UG 25000 pm and PG 20000 PM\' then calculate LPA and insrt \'3/2.4\' ')
            # print('In above example 3 represent UG salry in LPA and 2.4 represent PG salary in LPA ')
            print('Insert the proper salary in LPA............')
            st = input()
            salary_temp[i] = st

    return salary_temp


wb = xlrd.open_workbook(str(stname2))
sheet = wb.sheet_by_index(0)

ind_usn = []
ind_st = []
ind_deg = []
ind_major = []
ind_placed = []
ind_date = []

for i in range(sheet.nrows):
    for j in range(sheet.ncols):
        str1 = sheet.cell_value(i, j)
        str1 = str(str1).upper()
        if (str1.__contains__('USN')):
            str2 = sheet.cell_value(i + 1, j)
            str2 = str(str2).upper()
            if (str2.__contains__('1MS')):
                ind_usn.append([i, j])
                for k1 in range(sheet.ncols):
                    str3 = sheet.cell_value(i, k1)
                    str3 = str(str3).upper()
                    if (str3.__contains__('STUDENT')):
                        ind_st.append([i, k1])
                    if (str3.__contains__('DEGREE')):
                        ind_deg.append([i, k1])
                    if (str3.__contains__('MAJOR')):
                        ind_major.append([i, k1])
                    if (str3.__contains__('PLACED')):
                        ind_placed.append([i, k1])
                    if (str3.__contains__('DATE')):
                        ind_date.append([i, k1])

                break

# print(ind_usn,ind_st,ind_deg,ind_major,ind_placed,ind_date)
usn1 = []
student1 = []
deg1 = []
major1 = []
placed1 = []
date1 = []

for i in range(ind_usn[0][0] + 1, sheet.nrows):
    usn1.append(sheet.cell_value(i, ind_usn[0][1]))
    student1.append(sheet.cell_value(i, ind_st[0][1]))
    deg1.append(sheet.cell_value(i, ind_deg[0][1]))
    major1.append(sheet.cell_value(i, ind_major[0][1]))
    placed1.append(sheet.cell_value(i, ind_placed[0][1]))
    date1.append(sheet.cell_value(i, ind_date[0][1]))

for i in range(usn1.__len__()):

    str1 = usn1[i]
    str1 = str(str1).upper()
    if (usn1[i] == '' or (not (str1.__contains__('1MS')))):
        usn1[i] = 'NA11'
        student1[i] = 'NA11'
        deg1[i] = 'NA11'
        major1[i] = 'NA11'
        placed1[i] = 'NA11'
        date1[i] = 'NA11'

while 'NA11' in usn1:
    usn1.remove('NA11')
while 'NA11' in student1:
    student1.remove('NA11')
while 'NA11' in deg1:
    deg1.remove('NA11')
while 'NA11' in major1:
    major1.remove('NA11')
while 'NA11' in placed1:
    placed1.remove('NA11')
while 'NA11' in date1:
    date1.remove('NA11')
#
# print(usn1)
# print(student1)
# print(deg1)
# print(major1)
# print(placed1)
# print(date1)
#
# print(usn1.__len__())
# print(student1.__len__())
# print(deg1.__len__())
# print(major1.__len__())
# print(placed1.__len__())
# print(date1.__len__())

#####################################################


wb = xlrd.open_workbook(str(stname1))
sheet = wb.sheet_by_index(0)

ind_comp = []
temp = []

for i in range(sheet.nrows):
    for j in range(sheet.ncols):
        if (sheet.cell_value(i, j) == "Name of the Company"):
            # print(sheet.cell_value(i,j))
            temp.append(i)
            temp.append(j)
            ind_comp.append(temp)
            temp = []

        if (sheet.cell_value(i, j) == "Total Registered"):
            end_row = i

start_row = ind_comp[0][0] + 1

cse_ind = 0
date_ind = 0
salary_ind = 0
total_ind = 0
mcse_ind = 0
mcne_ind = 0

for i in range(sheet.ncols):
    str1 = sheet.cell_value(start_row, i)
    # print(str1)
    str1 = str(str1)
    str1 = str1.upper()

    if (str1.__contains__("M.TECH-CSE")):
        mcse_ind = i
    if (str1.__contains__("CSE") and (not str1.__contains__("TECH"))):
        cse_ind = i

    if (str1.__contains__("M.TECH-CNE")):
        mcne_ind = i
    if (str1.__contains__("TOTAL")):
        total_ind = i
    if (str1.__contains__("SALARY")):
        salary_ind = i
    if (str1.__contains__("DATE")):
        date_ind = i

# print(ind_comp)

comp_info = []
cse_info = []
mcse_info = []
mcne_info = []
total_info = []
salary_info = []
date_info = []

for i in range(start_row, end_row + 1):  # for rows
    str1 = sheet.cell_value(i, 1)
    str1 = str(str1).upper()
    if ((not str1.__contains__("COMPANY")) and str1 != ''):
        comp_info.append(sheet.cell_value(i, 1))
        if (isinstance(sheet.cell_value(i, cse_ind), str)):
            cse_info.append(sheet.cell_value(i, cse_ind))
        else:
            cse_info.append(int(sheet.cell_value(i, cse_ind)))
        if (isinstance(sheet.cell_value(i, mcse_ind), str)):
            mcse_info.append(sheet.cell_value(i, mcse_ind))
        else:
            mcse_info.append(int(sheet.cell_value(i, mcse_ind)))
        if (isinstance(sheet.cell_value(i, mcne_ind), str)):
            mcne_info.append(sheet.cell_value(i, mcne_ind))
        else:
            mcne_info.append(int(sheet.cell_value(i, mcne_ind)))
        total_info.append(sheet.cell_value(i, total_ind))
        salary_info.append(sheet.cell_value(i, salary_ind))
        date_info.append(sheet.cell_value(i, date_ind))

# 33333333333333333333333


for i in range(cse_info.__len__()):
    if ((cse_info[i] == 'NA' or cse_info[i] == 0.0) and (mcse_info[i] == 'NA' or mcse_info[i] == 0.0) and (
            mcne_info[i] == 'NA' or mcne_info[i] == 0.0)):
        comp_info[i] = 'NA11'
        cse_info[i] = 'NA11'
        mcse_info[i] = 'NA11'
        mcne_info[i] = 'NA11'
        total_info[i] = 'NA11'
        salary_info[i] = 'NA11'
        date_info[i] = 'NA11'

while 'NA11' in cse_info:
    cse_info.remove('NA11')
while 'NA11' in mcse_info:
    mcse_info.remove('NA11')
while 'NA11' in mcne_info:
    mcne_info.remove('NA11')
while 'NA11' in total_info:
    total_info.remove('NA11')
while 'NA11' in salary_info:
    salary_info.remove('NA11')
while 'NA11' in date_info:
    date_info.remove('NA11')
while 'NA11' in comp_info:
    comp_info.remove('NA11')

salary_info = salary_peocessing(salary_info)
salary_info = salary_check(salary_info)

t.start()


# print(comp_info)
# print(cse_info)
# print(mcse_info)
# print(mcne_info)
# print(total_info)
# print(salary_info)
# print(comp_info)


# t.start() #thread start for loading animation
# 333333333333333333333333333333333333333333333333333333333

# merging two table company and student to get salary column in student table



salary1 = ['0' for i in range(student1.__len__())]
comp1 = []
comp1 = ['NULL' for i in range(student1.__len__())]
usn_dict = {}

for i in range(student1.__len__()):

    st = placed1[i]
    st = str(st).upper()
    flag = 0

    for j in range(comp_info.__len__()):
        st1 = str(comp_info[j]).upper()
        st1 = st1.replace(" ", '')
        st = st.replace(" ", '')
        salary = 0
        if st.__contains__(st1) or st1.__contains__(st):

            # comp = comp_info[j]
            sal1 = str(salary_info[j]).split('/')
            if (sal1[0] != 'NULL'):
                if (usn1[1].__contains__('SCN') or usn1[i].__contains__('SCS') and sal1.__len__() == 2):

                    salary = float(sal1[1])
                else:
                    salary = float(sal1[0])


            salary1[i] = salary
            temp =comp_info[j]
            comp1[i] = temp
            # j = comp_info.__len__()

    #
    # if (salary == 0):
    #     nm = 'Enter salary of ' + placed1[i]
    #     salary = float(input(nm))
    #     salary1[i] = salary
    #     comp1[i] = placed1[i]
    #     comp_info.append(placed1[i])
    #     salary_info.append(salary)







        # else:
        #     comp = closest(comp_info, st)
        #     if (comp != ''):
        #
        #         for k1 in range(comp_info.__len__()):
        #             if (comp_info[k1] == comp):
        #                 salary = salary_info[k1]
        #                 comp1[k1] = st




###sorting all list based on usn



usn_ind = ['0' for i in range(usn1.__len__())]

for i in range(usn1.__len__()):
    st = usn1[i]
    if(st.__contains__('SCS') or st.__contains__('SCN')):
        usn_ind[i] = int((usn1[i][usn1[i].__len__()-2:]))
    else:
        usn_ind[i] = int((usn1[i][usn1[i].__len__() - 3:]))



usn1 = np.array(usn1)
comp1 = np.array(comp1)
salary1 = np.array(salary1)
student1 = np.array(student1)
deg1 = np.array(deg1)
major1 = np.array(major1)
placed1 = np.array(placed1)
date1 = np.array(date1)


usn_ind = np.array(usn_ind)
ind = usn_ind.argsort()

usn1 = usn1[ind]
comp1 = comp1[ind]
salary1 = salary1[ind]
student1 = student1[ind]
deg1 = deg1[ind]
major1 = major1[ind]
placed1 = placed1[ind]
date1 = date1[ind]


usn1 = np.array(usn1).tolist()
comp1 = np.array(comp1).tolist()
salary1 = np.array(salary1).tolist()
student1 = np.array(student1).tolist()
deg1 = np.array(deg1).tolist()
major1 = np.array(major1).tolist()
placed1 = np.array(placed1).tolist()
date1 = np.array(date1).tolist()


for i in range(salary1.__len__()):
    try:
        salary1[i] = float(salary1[i])
    except:
        salary1[i] = 0

#
# print(usn1)
# print(comp1)
# print(salary1)
# print(student1.__len__())
# print(deg1.__len__())
# print(major1.__len__())
# print(placed1.__len__())
# print(date1.__len__())

########## to remove duplicate usn
index_dict = {}

# salary1 = salary_check(salary1)

for i in range(salary1.__len__()):
    if not (isinstance(salary1[i],int) or isinstance(salary1[i],float) ):
        salary1[i] = 0

for i in range(usn1.__len__()):
    usn1[i] = str(usn1[i]).upper().replace(' ', '')
    usn_dict[usn1[i]] = []
    index_dict[usn1[i]]=[]

for i in range(usn1.__len__()):
    usn_dict[usn1[i]].append(salary1[i])
    index_dict[usn1[i]].append(i)

#33333333333333333333333333333333333333333333333333333333333
#student details with company name and salary(unique usn) PG

document = Document()

p = document.add_paragraph()
p.add_run('Ramaiah Institute Of Technology , Bangalore').bold = True

p1 = document.add_paragraph()
p1.add_run('The Department Of Training & Placement').bold = True

p1.add_run('------PG').bold = True

table = document.add_table(rows=1, cols=5, style='TableGrid')


hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Sl.no'
hdr_cells[1].text = 'USN'
hdr_cells[2].text = 'Student Name'
hdr_cells[3].text = 'Company'
hdr_cells[4].text = 'Salary'
# hdr_cells[5].text = 'Comp'

usn_list = []
off_pg_count =1
salary_temp = []




for i in range(student1.__len__()):
    usn = usn1[i].upper().replace(' ', '')

    if(usn.__contains__('SCS') or usn.__contains__('SCN') and usn!=''):
        if usn in usn_list:
            continue
        else:
            usn_list.append(usn)

        row_cells = table.add_row().cells
        row_cells[0].text = str(off_pg_count)
        off_pg_count += 1
        row_cells[1].text = usn1[i]
        row_cells[2].text = student1[i]

        max_sal = max(usn_dict[usn1[i]])
        l = usn_dict[usn1[i]]
        l1 = index_dict[usn1[i]]
        for k in range(l.__len__()):
            if (l[k] == max_sal):
                comp_name = comp1[l1[k]]

        row_cells[3].text = comp_name
        row_cells[4].text = str(max_sal) + ' LPA'
        salary_temp.append(max_sal)

    # print('with ' ,salary1[i] )
    # print('with max', max(usn_dict[usn1[i]]))

p2 = document.add_paragraph()
p2.add_run('\n')

pg_sal=[]
ug_sal=[]

# print(usn1)
# print(salary1)
for i in range(salary1.__len__()):
    cl = usn1[i]
    cl = str(cl).upper()
    if(cl.__contains__('SCS') or cl.__contains__('SCN')):
        pg_sal.append(salary1[i])
    else:
        ug_sal.append(salary1[i])

# print(pg_sal)
# print(ug_sal)
# exit()


# print(salary_temp)
# print(salary1)
# print(comp1)
# print(student1)
# print(usn1)
while 0 in pg_sal:
    pg_sal.remove(0)

while 0 in ug_sal:
    ug_sal.remove(0)


ls = pg_sal.__len__()
avg_sal = 0
num = 1
for i in range(int(ls - ls*0.9),int(ls - ls*0.1)):
    avg_sal+=pg_sal[i]
    num +=1
avg_sal = avg_sal/num


str1 = 'Minimum Salary : ' + str(min(pg_sal)) + ' LPA'
str1 = str(str1)
p2.add_run(str1)
p2.add_run('\n')

str1 = 'Maximum Salary : ' + str(max(pg_sal)) + ' LPA'
str1 = str(str1)
p2.add_run(str1)
p2.add_run('\n')



# print(avg_sal)

str1 = 'Average Salary : ' + str(round(avg_sal,2)) + ' LPA'
str1 = str(str1)
p2.add_run(str1)
p2.add_run('\n')


document.add_page_break()

# print(usn_list)

document.save('student_pg.docx')

#333333333333333333333333333333333333333333333333333333333333333333333333333333
# 33333333333333333333333333333333333333333333333333333333333
# student details with company name and salary(unique usn) UG



# document = Document('student_ug.docx')
document = Document()
p = document.add_paragraph()
p.add_run('Ramaiah Institute Of Technology , Bangalore').bold = True

p1 = document.add_paragraph()
p1.add_run('The Department Of Training & Placement').bold = True

p1.add_run('-----UG').bold = True

table = document.add_table(rows=1, cols=5, style='TableGrid')

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Sl.no'
hdr_cells[1].text = 'USN'
hdr_cells[2].text = 'Student Name'
hdr_cells[3].text = 'Company'
hdr_cells[4].text = 'Salary'
# hdr_cells[5].text = 'Comp'

usn_list = []
off_ug_count = 1
salary_temp=[]

for i in range(student1.__len__()):
    usn = usn1[i].upper().replace(' ', '')

    if not (usn.__contains__('SCS') or usn.__contains__('SCN') or usn==''):
        if usn in usn_list:
            continue
        else:
            usn_list.append(usn)

        row_cells = table.add_row().cells
        row_cells[0].text = str(off_ug_count)
        off_ug_count += 1
        row_cells[1].text = usn1[i]
        row_cells[2].text = student1[i]

        max_sal = max(usn_dict[usn1[i]])
        l = usn_dict[usn1[i]]
        l1 = index_dict[usn1[i]]
        for k in range(l.__len__()):
            if (l[k] == max_sal):
                comp_name = comp1[l1[k]]

        row_cells[3].text = comp_name
        row_cells[4].text = str(max_sal) + ' LPA'
        salary_temp.append(max_sal)

p2 = document.add_paragraph()
p2.add_run('\n')

str1 = 'Minimum Salary : ' + str(min(ug_sal)) + ' LPA'
str1 = str(str1)
p2.add_run(str1)
p2.add_run('\n')


str1 = 'Maximum Salary : ' + str(max(ug_sal)) + ' LPA'
str1 = str(str1)
p2.add_run(str1)
p2.add_run('\n')




ls = ug_sal.__len__()
avg_sal = 0
num = 1
for i in range(int(ls - ls * 0.9), int(ls - ls * 0.1)):
    avg_sal += ug_sal[i]
    num += 1
avg_sal = avg_sal / num

# print(avg_sal)




str1 = 'Average Salary : ' + str(round(avg_sal,2)) + ' LPA'
str1 = str(str1)
p2.add_run(str1)

    # print('with ' ,salary1[i] )
    # print('with max', max(usn_dict[usn1[i]]))

document.add_page_break()

# print(usn_list)

document.save('student_ug.docx')


# 333333333333333333333333333333333333333333333333333333333333333333333333333333

#students with 1,2,3,4,.....jobs PG


document = Document()

p = document.add_paragraph()
p.add_run('Ramaiah Institute Of Technology , Bangalore').bold = True

p1 = document.add_paragraph()
p1.add_run('The Department Of Training & Placement\n').bold = True

len = 0
key = list(usn_dict.keys())
for k in key:
    if(usn_dict[k].__len__() > len):
        len = usn_dict[k].__len__()
total =0

for k in range(len):
    usn_list = []
    count = 1
    p1 = document.add_paragraph()
    str1 = 'Students with ' + str(k+1) +' job'
    p1.add_run(str1).bold = True

    table = document.add_table(rows=1, cols=3, style='TableGrid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sl.no'
    hdr_cells[1].text = 'USN'
    hdr_cells[2].text = 'Student Name'

    for i in range(student1.__len__()):

        usn = usn1[i].upper().replace(' ','')

        if (usn.__contains__('SCS') or usn.__contains__('SCN') and usn != ''):
            if usn in usn_list:
                continue
            else:
                usn_list.append(usn)
            if (usn_dict[usn1[i]].__len__() == k + 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(count)
                count += 1
                row_cells[1].text = usn
                row_cells[2].text = student1[i]
                total+=1


    p1 = document.add_paragraph()
    p1.add_run('\n')
str1 = 'Total : '+str(total)
p1.add_run(str1)




document.add_page_break()

# print(usn_list)

document.save('no_of_offers_pg.docx')

#333333333333333333333333333333333333333333333333333333333333333333333333333

# fh = open('no_of_offers_ug.docx','w')
# fh.close()

#students with 1,2,3,4,.....jobs UG

document = Document()

p = document.add_paragraph()
p.add_run('Ramaiah Institute Of Technology , Bangalore').bold = True

p1 = document.add_paragraph()
p1.add_run('The Department Of Training & Placement\n').bold = True

len = 0
key = list(usn_dict.keys())
for k in key:
    if(usn_dict[k].__len__() > len):
        len = usn_dict[k].__len__()

total = 0

for k in range(len):
    usn_list = []
    count = 1
    p1 = document.add_paragraph()
    str1 = 'Students with ' + str(k+1) +' job'
    p1.add_run(str1).bold = True

    table = document.add_table(rows=1, cols=3, style='TableGrid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sl.no'
    hdr_cells[1].text = 'USN'
    hdr_cells[2].text = 'Student Name'


    for i in range(student1.__len__()):
        usn = usn1[i].upper().replace(' ','')

        if not(usn.__contains__('SCS') or usn.__contains__('SCN') or usn == ''):
            if usn in usn_list:
                continue
            else:
                usn_list.append(usn)
            if (usn_dict[usn1[i]].__len__() == k + 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(count)
                count += 1
                row_cells[1].text = usn
                row_cells[2].text = student1[i]
                total+=1

    p1 = document.add_paragraph()
    p1.add_run('\n')
str1 = 'Total : '+str(total)
p1.add_run(str1)


document.add_page_break()

# print(usn_list)

document.save('no_of_offers_ug.docx')

#333333333333333333333333333333333333333333333333333333333333333333333333333

#company wise salary  and count for PG
#
# fh = open('companywise_pg.docx','w')
# fh.close()

document = Document()

p = document.add_paragraph()
p.add_run('Ramaiah Institute Of Technology , Bangalore').bold = True

p1 = document.add_paragraph()
p1.add_run('The Department Of Training & Placement').bold = True

p1.add_run('-----PG').bold = True

table = document.add_table(rows=1, cols=5, style='TableGrid')

hdr_cells = table.rows[0].cells
hdr_cells[1].text = 'Company'
hdr_cells[2].text = 'Mtech - CSE'
hdr_cells[3].text = 'Mtech - CNE'
hdr_cells[4].text = 'Salary'
hdr_cells[0].text = 'Sl.no'

# hdr_cells[5].text = 'Comp'

usn_list = []
count = 1

pg_comp = []
pg_num = []

for i in range(comp_info.__len__()):

    str1 = str(comp_info[i])
    str1 = str1.upper()
    if mcse_info[i] == 'NA' or mcse_info[i] == '*':
        mcse_info[i] = 0
    if mcne_info[i] == 'NA' or mcne_info[i] == '*':
        mcne_info[i] = 0

    if(str1.__contains__('TOTAL')):
        continue

    if not(mcne_info[i]== 0 and mcse_info[i]== 0 ):

        row_cells = table.add_row().cells
        row_cells[0].text = str(count)
        count += 1
        row_cells[1].text = comp_info[i]
        row_cells[2].text = str(mcse_info[i])
        row_cells[3].text = str(mcne_info[i])
        sal1 = ''
        sal1 = salary_info[i]
        sal1 = str(sal1)
        if sal1.__contains__('/'):
            sal1 = sal1.split('/')
            row_cells[4].text = str(sal1[1]) + ' LPA'
        else:
            row_cells[4].text = str(sal1) + ' LPA'

        if not (isinstance(mcse_info[i],int) or isinstance(salary1[i],float)) :
            mcse_info[i] = 0
        if not (isinstance(mcne_info[i],int) or isinstance(salary1[i],float) ):
            mcne_info[i] = 0
        pg_comp.append(str(comp_info[i]))
        pg_num.append(mcne_info[i]+mcse_info[i])


p1 = document.add_paragraph()
p1.add_run('\n')
str1 = 'Total : '+str(count-1)
p1.add_run(str1)

table1 = document.add_table(rows=1, cols=4, style='TableGrid')

hdr_cells = table1.rows[0].cells
hdr_cells[1].text = 'Total Eligible'
hdr_cells[2].text = 'Total Placed'
hdr_cells[3].text = 'Unique Offers'

hdr_cells[0].text = 'Total Registered'

# 'Total Placed', 'Total Eligible', 'Total Registered'
row_cells = table1.add_row().cells
row_cells[0].text = str(mcse_info[mcse_info.__len__()-1] + mcne_info[mcne_info.__len__()-1])

row_cells[1].text = str(mcse_info[mcse_info.__len__()-2] + mcne_info[mcne_info.__len__()-2])
row_cells[2].text = str(mcse_info[mcse_info.__len__()-3] + mcne_info[mcne_info.__len__()-3])
row_cells[3].text = str(off_pg_count-1)




document.add_page_break()

# print(usn_list)



document.save('companywise_pg.docx')

# print(pg_comp)
# #plt.pie(cse_info1, labels=cse_comp,autopct='%1.1f%%', shadow=True, startangle=140)
# fig=plt.figure(figsize=(16,9))
# plt.bar(pg_comp, pg_num)
# plt.xticks(np.arange(0,pg_comp.__len__() ,1),rotation = 90)
# plt.yticks(np.arange(0,max(pg_num)+1 ,1))
# #plt.show()
# #fig.set_size_inches(18.5, 10.5)
# fig.savefig('pg_comp_graph.png', bbox_inches='tight')
# plt.close()

#plt.pie(cse_info1, labels=cse_comp,autopct='%1.1f%%', shadow=True, startangle=140)
fig=plt.figure(figsize=(16,9))
plt.bar(pg_comp, pg_num)
#plt.xticks(rotation=90)
plt.xticks(np.arange(0,pg_comp.__len__() ,1),rotation = 90)
plt.yticks(np.arange(0,max(pg_num)+1 ,1))
#plt.show()
#fig.set_size_inches(18.5, 10.5)
fig.savefig('pg_comp_graph.png', bbox_inches='tight')
plt.close()

#333333333333333333333333333333333333333333333333333333333333333333333333333



#company wise salary  and count for UG
#
# fh = open('companywise_ug.docx','w')
# fh.close()

document = Document()

p = document.add_paragraph()
p.add_run('Ramaiah Institute Of Technology , Bangalore').bold = True

p1 = document.add_paragraph()
p1.add_run('The Department Of Training & Placement').bold = True

p1.add_run('-----UG').bold = True

table = document.add_table(rows=1, cols=4, style='TableGrid')

hdr_cells = table.rows[0].cells
hdr_cells[1].text = 'Company'
hdr_cells[2].text = 'BE - CSE'
hdr_cells[3].text = 'Salary'
hdr_cells[0].text = 'Sl.no'

# hdr_cells[5].text = 'Comp'


usn_list = []
count = 1
ug_comp = []
ug_num = []

for i in range(comp_info.__len__()):

    str1 = str(comp_info[i])
    str1 = str1.upper()
    if cse_info[i] == 'NA' or cse_info[i] == '*':
        cse_info[i] = 0
    if cse_info[i] == 'NA' or cse_info[i] == '*':
        cse_info[i] = 0

    if(str1.__contains__('TOTAL')):
        continue

    if not(cse_info[i]== 0):

        row_cells = table.add_row().cells
        row_cells[0].text = str(count)
        count += 1
        row_cells[1].text = comp_info[i]
        row_cells[2].text = str(cse_info[i])
        # row_cells[3].text = str(mcne_info[i])
        sal1 = ''
        sal1 = salary_info[i]
        sal1 = str(sal1)
        if sal1.__contains__('/'):
            sal1 = sal1.split('/')
            row_cells[3].text = str(sal1[0]) + ' LPA'
        else:
            row_cells[3].text = str(sal1) + ' LPA'

        if not ( isinstance(cse_info[i],int) or isinstance(salary1[i],float) ):
            cse_info[i] = 0
        ug_comp.append(str(comp_info[i]))
        ug_num.append(cse_info[i])


p1 = document.add_paragraph()
p1.add_run('\n')
str1 = 'Total : '+str(count-1)
p1.add_run(str1)

table1 = document.add_table(rows=1, cols=4, style='TableGrid')

hdr_cells = table1.rows[0].cells
hdr_cells[1].text = 'Total Eligible'
hdr_cells[2].text = 'Total Placed'
hdr_cells[3].text = 'Unique Offers'

hdr_cells[0].text = 'Total Registered'

# 'Total Placed', 'Total Eligible', 'Total Registered'
row_cells = table1.add_row().cells
row_cells[0].text = str(cse_info[cse_info.__len__()-1])

row_cells[1].text = str(cse_info[cse_info.__len__()-2] )
row_cells[2].text = str(cse_info[cse_info.__len__()-3] )
row_cells[3].text = str(off_ug_count-1)




document.add_page_break()

# print(usn_list)

document.save('companywise_ug.docx')

# fig=plt.figure(figsize=(16,9))
# plt.bar(ug_comp, ug_num)
# plt.xticks(np.arange(0,ug_comp.__len__() ,1),rotation = 90)
# plt.yticks(np.arange(0,max(ug_num)+1 ,1))
# #plt.show()
# #fig.set_size_inches(18.5, 10.5)
# fig.savefig('ug_comp_graph.png', bbox_inches='tight')
# plt.close()


#plt.pie(cse_info1, labels=cse_comp,autopct='%1.1f%%', shadow=True, startangle=140)
fig=plt.figure(figsize=(16,9))
plt.bar(ug_comp, ug_num)
#plt.xticks(rotation=90)
plt.xticks(np.arange(0,ug_comp.__len__() ,1),rotation = 90)
plt.yticks(np.arange(0,max(ug_num)+1 ,1))
#plt.show()
#fig.set_size_inches(18.5, 10.5)
fig.savefig('ug_comp_graph.png', bbox_inches='tight')
plt.close()
#333333333333333333333333333333333333333333333333333333333333333333333333333







# print(list_count)
# print(list_count.__len__())
# print(list_comp.__len__())
# print(list_comp)
#
# print(usn1)
# print(placed1)
# print(salary_info)
# print(student1.__len__())
# print(comp_info)
# import matplotlib.pyplot as plt
#
# # Pie chart, where the slices will be ordered and plotted counter-clockwise:
# labels = list_comp
# sizes = list_count
# # explode = (0, 0.1, 0, 0)  # only "explode" the 2nd slice (i.e. 'Hogs')
#
#
# fig1, ax1 = plt.subplots()
# ax1.pie(sizes, labels=labels, autopct='%1.1f%%',
#         shadow=True, startangle=90)
# ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
#
# plt.show()
#
# plt.savefig('pie2.png')


# print('hey')
time.sleep(2)
done = True
